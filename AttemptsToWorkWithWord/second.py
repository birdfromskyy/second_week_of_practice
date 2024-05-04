import sys
import csv
import os
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from lxml import html
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
print(14)
def first_word(string):
    split_string = string.split()
    if split_string:
        return split_string[0]
def first_letter_of_second_word(string):
    split_string = string.split()
    if len(split_string) > 1:
        second_word = split_string[1]
        if second_word:  # Проверяем, что слово не пустое
            return second_word[0]

def first_letter_of_third_word(string):
    split_string = string.split()
    if len(split_string) > 2:
        third_word = split_string[2]
        if third_word:  # Проверяем, что слово не пустое
            return third_word[0]
print(32)
def convert_date_k(date_str):
    components = date_str.split("-")  # Разбиваем строку на компоненты
    if len(components) == 3:  # Проверяем, имеем ли мы только "год-день-месяц"
        year, day, month = components
    else:  # Если у нас есть все компоненты "год-день-месяц-день"
        year, day, month, _ = components
    return f"«{month}» {day} {year}"  # Возвращаем новую дату

def convert_date_p(date_str):
    components = date_str.split("-")  # Разбиваем строку на компоненты
    if len(components) == 3:  # Проверяем, имеем ли мы только "год-день-месяц"
        year, day, month = components
    else:  # Если у нас есть все компоненты "год-день-месяц-день"
        year, day, month, _ = components
    return f"{month}.{day}.{year}"  # Возвращаем новую дату

def parse_website(driver):
    tree = html.fromstring(driver.page_source)

    # Используйте XPath для извлечения данных, например:
    # data = tree.xpath("//div[@class='your-class']/text()")

    # Здесь предполагается, что вы извлекаете список данных из элементов с классом 'your-class'
    # Замените 'your-class' на соответствующий класс элемента
    # Верните список извлеченных данных
    data = ["Example data 1", "Example data 2", "Example data 3", "Example data 4", "Example data 5"]
    return data

args = sys.argv[1:]
if len(args) != 4:
    print("Необходимо передать четыре аргумента")
    sys.exit(1)
print(65)

auth_email = args[0]
auth_pass = args[1]
file_name_1 = args[2]
file_name_2 = args[3]

driver = webdriver.Chrome()
driver.get("http://localhost/SecondProject/login.php")

username = driver.find_element(By.NAME, "email")
username.send_keys(auth_email)
password = driver.find_element(By.NAME, "pass")
password.send_keys(auth_pass)
go = driver.find_element(By.NAME, "submit")
go.click()
go2 = driver.find_element(By.NAME, "apply_button")
go2.click()
print(83)


groupp = driver.find_element(By.XPATH, "/html/body/p[3]").text
vid = driver.find_element(By.XPATH, "/html/body/p[4]").text
type = driver.find_element(By.XPATH, "/html/body/p[5]").text
from_date = driver.find_element(By.XPATH, "/html/body/p[6]").text
to_date = driver.find_element(By.XPATH, "/html/body/p[7]").text
number_of_prikaz = driver.find_element(By.XPATH, "/html/body/p[8]").text
date = driver.find_element(By.XPATH, "/html/body/p[9]").text
direction = driver.find_element(By.XPATH, "/html/body/p[10]").text
kod = driver.find_element(By.XPATH, "/html/body/p[11]").text
vids = "учебной"

doc = Document('D:/xampp/htdocs/SecondProject/AttemptsToWorkWithWord/example2.docx')

print(99)
file_path_1 = os.path.join('D:\\xampp\\htdocs\\SecondProject\\AttemptsToWorkWithWord\\', file_name_1)
file_path_2 = os.path.join('D:\\xampp\\htdocs\\SecondProject\\AttemptsToWorkWithWord\\', file_name_2)
if os.path.exists(file_path_1):
        # Открываем файл для чтения
        with open(file_path_1, newline='', encoding='utf-8') as csvfile:
            csv_reader = csv.reader(csvfile)

            fios1 = []

            # Читаем каждую строку CSV-файла
            for row in csv_reader:
                # Если строка не пустая, добавляем значение из первого столбца в список fios1
                if row:
                    fios1.append(row[0])
print(114)
for table in doc.tables:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text.strip() == "FIO1":
                current_row_index = i
                current_column_index = j
                for k, fio_value in enumerate(fios1):
                    table.cell(current_row_index + k, current_column_index).text = fio_value
                break  # Прерываем цикл, чтобы заменить только один столбец

print(125)
print(len(fios1))
if os.path.exists(file_path_2):
    # Открываем файл для чтения
    with open(file_path_2, newline='', encoding='utf-8') as csvfile:
        csv_reader = csv.reader(csvfile)

        fios2 = []

        # Читаем каждую строку CSV-файла
        for row in csv_reader:
            # Если строка не пустая, добавляем значение из первого столбца в список fios2
            if row:
                fios2.append(row[0])


for table in doc.tables:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text.strip() == "FIO2":
                current_row_index = i
                current_column_index = j
                for k, fio_value in enumerate(fios2):
                    table.cell(current_row_index + k, current_column_index).text = fio_value
                break  # Прерываем цикл, чтобы заменить только один столбец

for table in doc.tables:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text.strip() == "UGRASU":  # Находим метку "UGRASU"
                current_row_index = i
                current_column_index = j
                # Вставляем "ЮГУ" от 0 до len(fios1)
                for k in range(len(fios2)):
                    table.cell(current_row_index + k, current_column_index).text = " ЮГУ"
                break  # Прерываем цикл, чтобы заменить только один столбец

for table in doc.tables:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text.strip() == "HANT":  # Находим метку "UGRASU"
                current_row_index = i
                current_column_index = j
                # Вставляем "ЮГУ" от 0 до len(fios1)
                for k in range(len(fios2)):
                    table.cell(current_row_index + k, current_column_index).text = " г.Ханты-Мансийск"
                break  # Прерываем цикл, чтобы заменить только один столбец

for table in doc.tables:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text.strip() == "RUK":  # Находим метку "UGRASU"
                current_row_index = i
                current_column_index = j
                # Вставляем "ЮГУ" от 0 до len(fios1)
                for k in range(len(fios2)):
                    table.cell(current_row_index + k, current_column_index).text = " Змеев Д.О."
                break  # Прерываем цикл, чтобы заменить только один столбец


def format_table_cells(table):
    """
    Форматирует текст в ячейках таблицы.

    :param table: Объект таблицы для форматирования.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Устанавливаем выравнивание текста в ячейке на центр
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in paragraph.runs:
                    # Устанавливаем шрифт на 'Times New Roman'
                    run.font.name = 'Times New Roman'
                    # Устанавливаем размер шрифта на 11 пунктов
                    run.font.size = Pt(11)
countt = len(fios2)

def format_tables_in_document(doc):
    """
    Форматирует текст в ячейках всех таблиц в документе.

    :param doc: Объект документа для форматирования.
    """
    # Проходим по всем таблицам в документе
    for table in doc.tables:
        format_table_cells(table)

def replace_text_in_document(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)
print(224)
counttz2 = len(fios2)
counttz1 = len(fios1)
replace_text_in_document(doc, '{{countt}}', str(counttz2))
replace_text_in_document(doc, '{{counttz}}', str(counttz1))
replace_text_in_document(doc, '{{groupp}}', groupp)
replace_text_in_document(doc, '{{direction}}', direction)
replace_text_in_document(doc, '{{kod}}', kod)
format_tables_in_document(doc)
print(228)
doc.save('second.docx')

extracted_data = parse_website(driver)
driver.quit()