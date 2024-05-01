# from docx import Document
# def replace_text_in_document(doc, old_text, new_text):
#     for paragraph in doc.paragraphs:
#         if old_text in paragraph.text:
#             paragraph.text = paragraph.text.replace(old_text, new_text)
# doc = Document('D:/xampp/htdocs/SecondProject/AttemptsToWorkWithWord/exp.docx')
# replace_text_in_document(doc, '{{FIO}}', 'Ножников Геннадий Александрович')
# replace_text_in_document(doc, '{{AGE}}', '20')
# doc.save('output.docx')
# print('Finish')
import sys
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By

from lxml import html
from docx import Document


def replace_text_in_document(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

# Функция для парсинга данных с использованием XPath
def parse_website(driver):
    tree = html.fromstring(driver.page_source)
    
    # Используйте XPath для извлечения данных, например:
    # data = tree.xpath("//div[@class='your-class']/text()")
    
    # Здесь предполагается, что вы извлекаете список данных из элементов с классом 'your-class'
    # Замените 'your-class' на соответствующий класс элемента
    # Верните список извлеченных данных
    data = ["Example data 1", "Example data 2", "Example data 3", "Example data 4", "Example data 5"]
    return data

# Функция для создания документа DOCX и добавления извлеченных строк
def create_docx(data):
    doc = Document()
    for item in data:
        doc.add_paragraph(item)
    doc.save("exp.docx")

args = sys.argv[1:]
auth_email = args[0]
auth_pass = args[1]


driver = webdriver.Chrome()
# Создание экземпляра драйвера Selenium (Chrome)
#chrome_options = Options()

# Установка безголового режима
#chrome_options.add_argument("--headless")

# Создание экземпляра драйвера с указанием опций
#driver = webdriver.Chrome(options=chrome_options)

# Загрузка страницы для авторизации
driver.get("http://localhost/SecondProject/login.php")


#Ввод логина
username = driver.find_element(By.NAME, "email")
username.send_keys(auth_email)
#Ввод пароля
password = driver.find_element(By.NAME, "pass")
password.send_keys(auth_pass)
#Авторизация
go = driver.find_element(By.NAME, "submit")
go.click()
#Личные данные
email = driver.find_element(By.XPATH, "/html/body/div/div/div[1]/span").text
course = driver.find_element(By.XPATH, "/html/body/div/div/div[2]/span").text
groupp = driver.find_element(By.XPATH, "/html/body/div/div/div[3]/span").text
name = driver.find_element(By.XPATH, "/html/body/div/div/div[4]/span").text
surname = driver.find_element(By.XPATH, "/html/body/div/div/div[5]/span").text
patronymic = driver.find_element(By.XPATH, "/html/body/div/div/div[6]/span").text
institute = driver.find_element(By.XPATH, "/html/body/div/div/div[7]/span").text
direction = driver.find_element(By.XPATH, "/html/body/div/div/div[8]/span").text
#Данные о взаимодействии
payment = driver.find_element(By.XPATH, "/html/body/div/div/div[9]/span").text
spravlyalsya = driver.find_element(By.XPATH, "/html/body/div/div/div[10]/span").text
kachestva = driver.find_element(By.XPATH, "/html/body/div/div/div[11]/span").text
amount_vipolneniya = driver.find_element(By.XPATH, "/html/body/div/div/div[12]/span").text
zamechaniya = driver.find_element(By.XPATH, "/html/body/div/div/div[13]/span").text
ocenka = driver.find_element(By.XPATH, "/html/body/div/div/div[14]/span").text
#Данные о практике
vid = driver.find_element(By.XPATH, "/html/body/div/div/div[15]/span").text
type = driver.find_element(By.XPATH, "/html/body/div/div/div[16]/span").text
name_org = driver.find_element(By.XPATH, "/html/body/div/div/div[17]/span").text
place_practic = driver.find_element(By.XPATH, "/html/body/div/div/div[18]/span").text
address_organization = driver.find_element(By.XPATH, "/html/body/div/div/div[19]/span").text
fio_director_of_company = driver.find_element(By.XPATH, "/html/body/div/div/div[20]/span").text
post_director_of_company = driver.find_element(By.XPATH, "/html/body/div/div/div[21]/span").text
fio_director_of_ugrasu = driver.find_element(By.XPATH, "/html/body/div/div/div[22]/span").text
post_director_of_ugrasu = driver.find_element(By.XPATH, "/html/body/div/div/div[23]/span").text
fio_director_of_organization = driver.find_element(By.XPATH, "/html/body/div/div/div[24]/span").text
post_director_of_organization = driver.find_element(By.XPATH, "/html/body/div/div/div[25]/span").text
from_date = driver.find_element(By.XPATH, "/html/body/div/div/div[26]/span").text
to_date = driver.find_element(By.XPATH, "/html/body/div/div/div[27]/span").text
number_of_prikaz = driver.find_element(By.XPATH, "/html/body/div/div/div[28]/span").text
date = driver.find_element(By.XPATH, "/html/body/div/div/div[29]/span").text
#Задачи
task1 = driver.find_element(By.XPATH, "/html/body/div/div/div[30]/span").text
date1 = driver.find_element(By.XPATH, "/html/body/div/div/div[31]/span").text
task2 = driver.find_element(By.XPATH, "/html/body/div/div/div[32]/span").text
date2 = driver.find_element(By.XPATH, "/html/body/div/div/div[33]/span").text
task3 = driver.find_element(By.XPATH, "/html/body/div/div/div[34]/span").text
date3 = driver.find_element(By.XPATH, "/html/body/div/div/div[35]/span").text
task4 = driver.find_element(By.XPATH, "/html/body/div/div/div[36]/span").text
date4 = driver.find_element(By.XPATH, "/html/body/div/div/div[37]/span").text
task5 = driver.find_element(By.XPATH, "/html/body/div/div/div[38]/span").text
date5 = driver.find_element(By.XPATH, "/html/body/div/div/div[39]/span").text
task6 = driver.find_element(By.XPATH, "/html/body/div/div/div[40]/span").text
date6 = driver.find_element(By.XPATH, "/html/body/div/div/div[41]/span").text
task7 = driver.find_element(By.XPATH, "/html/body/div/div/div[42]/span").text
date7 = driver.find_element(By.XPATH, "/html/body/div/div/div[43]/span").text
task8 = driver.find_element(By.XPATH, "/html/body/div/div/div[44]/span").text
date8 = driver.find_element(By.XPATH, "/html/body/div/div/div[45]/span").text
task9 = driver.find_element(By.XPATH, "/html/body/div/div/div[46]/span").text
date9 = driver.find_element(By.XPATH, "/html/body/div/div/div[47]/span").text
task10 = driver.find_element(By.XPATH, "/html/body/div/div/div[48]/span").text
date10 = driver.find_element(By.XPATH, "/html/body/div/div/div[49]/span").text
task11 = driver.find_element(By.XPATH, "/html/body/div/div/div[50]/span").text
date11 = driver.find_element(By.XPATH, "/html/body/div/div/div[51]/span").text
task12 = driver.find_element(By.XPATH, "/html/body/div/div/div[52]/span").text
date12 = driver.find_element(By.XPATH, "/html/body/div/div/div[53]/span").text
task13 = driver.find_element(By.XPATH, "/html/body/div/div/div[54]/span").text
date13 = driver.find_element(By.XPATH, "/html/body/div/div/div[55]/span").text
task14 = driver.find_element(By.XPATH, "/html/body/div/div/div[56]/span").text
date14 = driver.find_element(By.XPATH, "/html/body/div/div/div[57]/span").text
task15 = driver.find_element(By.XPATH, "/html/body/div/div/div[58]/span").text
date15 = driver.find_element(By.XPATH, "/html/body/div/div/div[59]/span").text
task16 = driver.find_element(By.XPATH, "/html/body/div/div/div[60]/span").text
date16 = driver.find_element(By.XPATH, "/html/body/div/div/div[61]/span").text
task17 = driver.find_element(By.XPATH, "/html/body/div/div/div[62]/span").text
date17 = driver.find_element(By.XPATH, "/html/body/div/div/div[63]/span").text
task18 = driver.find_element(By.XPATH, "/html/body/div/div/div[64]/span").text
date18 = driver.find_element(By.XPATH, "/html/body/div/div/div[65]/span").text
task19 = driver.find_element(By.XPATH, "/html/body/div/div/div[66]/span").text
date19 = driver.find_element(By.XPATH, "/html/body/div/div/div[67]/span").text
task20 = driver.find_element(By.XPATH, "/html/body/div/div/div[68]/span").text
date20 = driver.find_element(By.XPATH, "/html/body/div/div/div[69]/span").text
task21 = driver.find_element(By.XPATH, "/html/body/div/div/div[70]/span").text
date21 = driver.find_element(By.XPATH, "/html/body/div/div/div[71]/span").text



doc = Document('D:/xampp/htdocs/SecondProject/AttemptsToWorkWithWord/example.docx')
replace_text_in_document
replace_text_in_document(doc, '{{email}}', email)
replace_text_in_document(doc, '{{course}}', course)
replace_text_in_document(doc, '{{groupp}}', groupp)
replace_text_in_document(doc, '{{name}}', name)
replace_text_in_document(doc, '{{surname}}', surname)
replace_text_in_document(doc, '{{patronymic}}', patronymic)
replace_text_in_document(doc, '{{institute}}', institute)
replace_text_in_document(doc, '{{direction}}', direction)
replace_text_in_document(doc, '{{payment}}', payment)
replace_text_in_document(doc, '{{spravlyalsya}}', spravlyalsya)
replace_text_in_document(doc, '{{kachestva}}', kachestva)
replace_text_in_document(doc, '{{amount_vipolneniya}}',amount_vipolneniya)
replace_text_in_document(doc, '{{zamechaniya}}', zamechaniya)
replace_text_in_document(doc, '{{ocenka}}', ocenka)
replace_text_in_document(doc, '{{vid}}', vid)
replace_text_in_document(doc, '{{type}}', type)
replace_text_in_document(doc, '{{name_org}}', name_org)
replace_text_in_document(doc, '{{place_practic}}', place_practic)
replace_text_in_document(doc, '{{address_organization}}', address_organization)
replace_text_in_document(doc, '{{fio_director_of_company}}', fio_director_of_company)
replace_text_in_document(doc, '{{post_director_of_company}}', post_director_of_company)
replace_text_in_document(doc, '{{fio_director_of_ugrasu}}', fio_director_of_ugrasu)
replace_text_in_document(doc, '{{post_director_of_ugrasu}}', post_director_of_ugrasu)
replace_text_in_document(doc, '{{fio_director_of_organization}}', fio_director_of_organization)
replace_text_in_document(doc, '{{post_director_of_organization}}', post_director_of_organization)
replace_text_in_document(doc, '{{from_date}}', from_date)
replace_text_in_document(doc, '{{to_date}}', to_date)
replace_text_in_document(doc, '{{number_of_prikaz}}', number_of_prikaz)
replace_text_in_document(doc, '{{date}}', date)
replace_text_in_document(doc, '{{task1}}', task1)
replace_text_in_document(doc, '{{date1}}', date1)
replace_text_in_document(doc, '{{task2}}', task2)
replace_text_in_document(doc, '{{date2}}', date2)
replace_text_in_document(doc, '{{task3}}', task3)
replace_text_in_document(doc, '{{date3}}', date3)
replace_text_in_document(doc, '{{task4}}', task4)
replace_text_in_document(doc, '{{date4}}', date4)
replace_text_in_document(doc, '{{task5}}', task5)
replace_text_in_document(doc, '{{date5}}', date5)
replace_text_in_document(doc, '{{task6}}', task6)
replace_text_in_document(doc, '{{date6}}', date6)
replace_text_in_document(doc, '{{task7}}', task7)
replace_text_in_document(doc, '{{date7}}', date7)
replace_text_in_document(doc, '{{task8}}', task8)
replace_text_in_document(doc, '{{date8}}', date8)
replace_text_in_document(doc, '{{task9}}', task9)
replace_text_in_document(doc, '{{date9}}', date9)
replace_text_in_document(doc, '{{task10}}', task10)
replace_text_in_document(doc, '{{date10}}', date10)
replace_text_in_document(doc, '{{task11}}', task11)
replace_text_in_document(doc, '{{date11}}', date11)
replace_text_in_document(doc, '{{task12}}', task12)
replace_text_in_document(doc, '{{date12}}', date12)
replace_text_in_document(doc, '{{task13}}', task13)
replace_text_in_document(doc, '{{date13}}', date13)
replace_text_in_document(doc, '{{task14}}', task14)
replace_text_in_document(doc, '{{date14}}', date14)
replace_text_in_document(doc, '{{task15}}', task15)
replace_text_in_document(doc, '{{date15}}', date15)
replace_text_in_document(doc, '{{task16}}', task16)
replace_text_in_document(doc, '{{date16}}', date16)
replace_text_in_document(doc, '{{task17}}', task17)
replace_text_in_document(doc, '{{date17}}', date17)
replace_text_in_document(doc, '{{task18}}', task18)
replace_text_in_document(doc, '{{date18}}', date18)
replace_text_in_document(doc, '{{task19}}', task19)
replace_text_in_document(doc, '{{date19}}', date19)
replace_text_in_document(doc, '{{task20}}', task20)
replace_text_in_document(doc, '{{date20}}', date20)
replace_text_in_document(doc, '{{task21}}', task21)
replace_text_in_document(doc, '{{date21}}', date21)
replace_text_in_document(doc, '0000-00-00',"")
tasks = []
tasks.append(task1)
tasks.append(task2)
tasks.append(task3)
tasks.append(task4)
tasks.append(task5)
tasks.append(task6)
tasks.append(task7)
tasks.append(task8)
tasks.append(task9)
tasks.append(task10)
tasks.append(task11)
tasks.append(task12)
tasks.append(task13)
tasks.append(task14)
tasks.append(task15)
tasks.append(task16)
tasks.append(task17)
tasks.append(task18)
tasks.append(task19)
tasks.append(task20)
tasks.append(task21)


def insert_tasks_between_text(document_path, start_text, end_text, *tasks):
    # Открываем существующий документ
    doc = Document(document_path)

    # Поиск начала и конца фрагмента для замены
    start_index = None
    end_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if start_text in paragraph.text:
            start_index = i
        elif end_text in paragraph.text:
            end_index = i
            break

    # Проверка наличия меток текста в документе
    if start_index is None or end_index is None:
        print("Метки текста не найдены в документе.")
        return

    # Удаление содержимого между метками
    for i in range(start_index + 1, end_index):
        doc.paragraphs[start_index + 1].clear()

    # Вставка задач вместо удаленного содержимого
    for i, task_text in enumerate(tasks, start=1):
        doc.paragraphs[start_index + 1].insert_paragraph_before(f"{i}) {task_text}")

    # Сохранение изменений
    doc.save(document_path)


# Пример использования функции
start_text = "start_text"
end_text = "end_text"
insert_tasks_between_text(doc, start_text, end_text, *tasks)

doc.save('output.docx')


# Ждем, пока страница загрузится
driver.implicitly_wait(10)

# Получаем данные с помощью XPath
extracted_data = parse_website(driver)

# Создаем и заполняем документ DOCX
# create_docx(extracted_data)
# Закрытие браузера
driver.quit()